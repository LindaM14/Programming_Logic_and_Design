from openpyxl import Workbook
import requests
import docx


# API URL
word_bank = 'http://api.worldbank.org/v2/country?format=json&per_page=400'
response = requests.get(word_bank).json()

# Entering the page/index of the dictionary
country_info_list = response[1]

# creating the word docx sheet
document = docx.Document()
document.add_paragraph('Countries and their Capital Cities', 'Title')

# creating the excel sheet
capital_cities_workbook = Workbook()
worksheet = capital_cities_workbook.active
worksheet.title = 'Capital Cities'
worksheet.cell(1, 1,)
worksheet.cell(1, 2,)

# for loop to make data go into excel and word sheet in dictionary Name and CapitalCity
for country_info_dictionary, country_info_list in enumerate(country_info_list):
    worksheet.cell(country_info_dictionary +2, 1, country_info_list['name'])
    worksheet.cell(country_info_dictionary + 2, 2, country_info_list['capitalCity'])
    document.add_paragraph(f'The capital of ' + country_info_list['name'] + ' is ' + country_info_list['capitalCity'])

# save documents - when print will copy/data transfer to sheets
capital_cities_workbook.save('capital_cities.xls')
document.save('Capital_Cities.docx')